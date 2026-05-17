"""
Generador de documentos DOCX con estilos corporativos.
Crea archivos Word con logo, header/footer, estilos de texto, 
bloques de código con fondo oscuro y tablas formateadas.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from typing import Dict, Any, List, Optional
import io


class DOCXGenerator:
    def __init__(self, logo_path: Optional[str] = None):
        self.logo_path = logo_path
        self.document = None

    def setup_styles(self, doc: Document):
        """Configura los estilos del documento según la plantilla corporativa."""
        # Estilo Normal
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Estilo Heading 1 - Margen amplio para separar del contenido
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        h1.paragraph_format.space_after = Pt(18)
        h1.paragraph_format.space_before = Pt(12)

        # Estilo Heading 2
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        h2.paragraph_format.space_after = Pt(14)
        h2.paragraph_format.space_before = Pt(8)

        # Estilo Heading 3
        h3 = doc.styles['Heading 3']
        h3.font.name = 'Calibri'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        h3.paragraph_format.space_after = Pt(12)
        h3.paragraph_format.space_before = Pt(6)

        # Crear estilo de código
        try:
            code_style = doc.styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
        except:
            code_style = doc.styles['CodeStyle']
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)

    def add_logo(self, doc: Document):
        """Agrega el logo centrado en la parte superior de la primera página."""
        if self.logo_path and os.path.exists(self.logo_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(self.logo_path, width=Inches(2.0))
            doc.add_paragraph()  # Espacio después del logo

    def add_title(self, doc: Document, metadata: Dict[str, Any]):
        """Agrega el título principal del documento con estilo List Paragraph."""
        dmnd = metadata.get('numero_demanda', '')
        titulo = metadata.get('titulo', '')

        # Evitar duplicar "DMND" si el usuario ya lo incluyó en el título
        if dmnd and dmnd.strip() in titulo:
            full_title = titulo
        elif dmnd and dmnd.strip():
            full_title = f"DMND{dmnd} - {titulo}"
        else:
            full_title = titulo

        # Usar estilo List Paragraph como en el documento original
        paragraph = doc.add_paragraph(style='List Paragraph')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(full_title)
        run.bold = True
        run.font.size = Pt(36)
        run.font.name = 'Calibri'
        doc.add_paragraph()

    def add_metadata_section(self, doc: Document, metadata: Dict[str, Any]):
        """Agrega la sección de metadatos debajo del título."""
        fields = [
            ("Información Adicional", None),
            ("Número de llamado", metadata.get('numero_demanda', '')),
            ("Fecha", metadata.get('fecha', '')),
            ("Ciclo", metadata.get('ciclo', '')),
            ("Sistema/Módulo", metadata.get('sistema', 'COMEX')),
            ("Versión", metadata.get('version', '1.0')),
            ("Autor", metadata.get('autor', ''))
        ]

        for label, value in fields:
            if value is None:
                # Es un título de grupo
                p = doc.add_paragraph()
                run = p.add_run(label)
                run.bold = True
                run.font.size = Pt(12)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: {value}")
                run.font.size = Pt(11)

        doc.add_paragraph()  # Espacio
        # Salto de página después de la portada para que contenido inicie en página 2
        doc.add_page_break()

    def add_toc(self, doc: Document, sections: List[Dict[str, Any]]):
        """Agrega un índice de contenido (TOC) nativo de Word."""
        from docx.oxml import OxmlElement
        
        # Título "Índice"
        heading = doc.add_heading("Índice", level=1)
        
        # Crear párrafo para el campo TOC
        paragraph = doc.add_paragraph()
        
        # Inicio del campo
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        # Instrucción del campo TOC
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run2._r.append(instrText)
        
        # Separador
        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar2)
        
        # Texto placeholder
        run4 = paragraph.add_run("Haga clic derecho y seleccione 'Actualizar campo' para ver el indice")
        run4.font.italic = True
        run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        # Fin del campo
        run5 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar3)
        
        # Agregar setting para actualizar campos al abrir
        settings = doc.settings
        settings_element = settings.element
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings_element.append(update_fields)
        
        doc.add_page_break()

    def add_code_block(self, doc: Document, code: str, language: str = "sql"):
        """Agrega un bloque de código con fondo oscuro y tipografía consola."""
        # Crear tabla de 1x1 para simular bloque de código con fondo
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Estilo de la celda (fondo oscuro)
        cell = table.cell(0, 0)
        shading_elm = parse_xml(r'<w:shd {} w:fill="1E1E1E"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Agregar código
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)

        # Margen interno
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.space_after = Pt(0)
        paragraph.space_before = Pt(0)

        doc.add_paragraph()  # Espacio después del bloque

    def add_table(self, doc: Document, headers: List[str], rows: List[List[str]]):
        """Agrega una tabla con estilos corporativos. Filtra filas completamente vacías."""
        if not headers:
            return
        
        # Filtrar filas completamente vacías
        filtered_rows = []
        for row in rows:
            # Verificar si la fila tiene al menos un valor no vacío
            has_content = any(str(cell).strip() for cell in row)
            if has_content:
                filtered_rows.append(row)
        
        if not filtered_rows:
            return

        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            # Fondo gris para header
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            # Texto en negrita
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Data rows (filtradas)
        for row_data in filtered_rows:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell_text)
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()  # Espacio después de tabla

    def _add_formatted_paragraph(self, doc: Document, text: str, style: str = 'Normal'):
        """Agrega un párrafo interpretando markdown básico (bold, inline code)."""
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.space_before = Pt(4)

        import re
        # Patrón: **bold**, `code`, o texto normal
        pattern = r'(\*\*(.+?)\*\*|`(.+?)`|(.+?))(?=\*\*|$|`)'
        # Simpler: split by ** and `
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)

        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(11)
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
                # Shading
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
                run.font.highlight_color = None
            else:
                # Normal text
                run = p.add_run(part)
                run.font.size = Pt(11)

    def add_content_section(self, doc: Document, title: str, text: str, table_data: Optional[Dict] = None):
        """Agrega una sección completa: título, texto y tabla opcional."""
        # Título
        if title:
            p = doc.add_heading(title, level=1)
            # Agregar espacio después del título
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(6)

        # Procesar texto (incluyendo tablas markdown, código, headings, listas)
        if text:
            self._process_text_with_tables(doc, text)

        # Tabla del data editor (si existe y está habilitada)
        if table_data and table_data.get('headers'):
            self.add_table(doc, table_data['headers'], table_data.get('rows', []))

    def _flush_list(self, doc: Document, items: list):
        """Agrega items acumulados como lista con viñetas."""
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            self._apply_inline_formatting(p, item)
            p.paragraph_format.space_after = Pt(4)

    def _apply_inline_formatting(self, paragraph, text: str):
        """Aplica bold e inline code a un párrafo existente."""
        import re
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(11)
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                run = paragraph.add_run(part)
                run.font.size = Pt(11)

    def add_header_footer(self, doc: Document, metadata: Dict[str, Any]):
        """Agrega header y footer corporativos a todas las páginas."""
        section = doc.sections[0]
        
        # HEADER: Tabla 1x2 - "COMEX" | "DOCUMENTO FUNCIONAL"
        header = section.header
        header.is_linked_to_previous = False
        
        # Limpiar header existente
        for p in header.paragraphs:
            p.clear()
        
        # Agregar tabla al header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Color corporativo extraído del documento original: ED7D31 (naranja)
        HEADER_BG_COLOR = "ED7D31"
        HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)  # Blanco
        
        # Celda izquierda: COMEX
        left_cell = header_table.cell(0, 0)
        # Fondo naranja
        shading_left = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), HEADER_BG_COLOR))
        left_cell._tc.get_or_add_tcPr().append(shading_left)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run("COMEX")
        left_run.bold = True
        left_run.font.size = Pt(10)
        left_run.font.color.rgb = HEADER_TEXT_COLOR
        
        # Celda derecha: DOCUMENTO FUNCIONAL
        right_cell = header_table.cell(0, 1)
        # Fondo naranja
        shading_right = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), HEADER_BG_COLOR))
        right_cell._tc.get_or_add_tcPr().append(shading_right)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run("DOCUMENTO FUNCIONAL")
        right_run.bold = True
        right_run.font.size = Pt(10)
        right_run.font.color.rgb = HEADER_TEXT_COLOR
        
        # FOOTER: Tabla 1x3 - Info documento | Fecha | Página
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Limpiar footer existente
        for p in footer.paragraphs:
            p.clear()
        
        # Agregar tabla al footer
        footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        dmnd = metadata.get('numero_demanda', '')
        titulo = metadata.get('titulo', '')
        fecha = metadata.get('fecha', '')
        
        # Celda izquierda: DMND + Título
        f_left = footer_table.cell(0, 0)
        f_left_para = f_left.paragraphs[0]
        f_left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        f_left_run = f_left_para.add_run(f"DMND{dmnd}")
        f_left_run.font.size = Pt(8)
        f_left_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        if titulo:
            f_left_para.add_run(f" - {titulo[:30]}").font.size = Pt(8)
            f_left_para.runs[1].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # Celda centro: Fecha
        f_center = footer_table.cell(0, 1)
        f_center_para = f_center.paragraphs[0]
        f_center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_center_run = f_center_para.add_run(f"{fecha}")
        f_center_run.font.size = Pt(8)
        f_center_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # Celda derecha: Número de página (campo)
        f_right = footer_table.cell(0, 2)
        f_right_para = f_right.paragraphs[0]
        f_right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Agregar campo de número de página
        from docx.oxml import OxmlElement
        run = f_right_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def add_page_border(self, doc: Document):
        """Agrega borde de página al documento."""
        section = doc.sections[0]
        sectPr = section._sectPr
        
        # Crear elemento pgBorders
        pgBorders = parse_xml(
            r'<w:pgBorders {} w:offsetFrom="page">'
            r'<w:top w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
            r'<w:left w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
            r'<w:right w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
            r'</w:pgBorders>'.format(nsdecls('w'))
        )
        sectPr.append(pgBorders)

    def _parse_markdown_table(self, lines: List[str], start_idx: int) -> tuple:
        """Parsea una tabla markdown y retorna (headers, rows, end_idx)."""
        table_lines = []
        i = start_idx
        
        # Recolectar líneas de la tabla
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        
        if len(table_lines) < 2:
            return None, None, start_idx
        
        # Parsear header (primera línea)
        header_line = table_lines[0]
        headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
        
        # Ignorar línea separadora (segunda línea)
        # Parsear filas de datos (desde tercera línea)
        rows = []
        for line in table_lines[2:]:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
        
        return headers, rows, i

    def _process_text_with_tables(self, doc: Document, text: str):
        """Procesa texto markdown detectando tablas y convirtiéndolas a tablas reales."""
        lines = text.split('\n')
        i = 0
        in_code_block = False
        code_buffer = []
        language = "sql"
        list_buffer = []
        in_list = False
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Detectar inicio/fin de bloque de código
            if stripped.startswith('```'):
                if in_code_block:
                    if code_buffer:
                        self.add_code_block(doc, '\n'.join(code_buffer), language)
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                    lang = stripped[3:].strip()
                    if lang:
                        language = lang
                i += 1
                continue
            
            if in_code_block:
                code_buffer.append(line)
                i += 1
                continue
            
            # Detectar tabla markdown
            if stripped.startswith('|') and not in_list:
                headers, rows, end_idx = self._parse_markdown_table(lines, i)
                if headers and rows is not None:
                    # Flush cualquier lista pendiente
                    if in_list and list_buffer:
                        self._flush_list(doc, list_buffer)
                        list_buffer = []
                        in_list = False
                    # Agregar tabla real
                    self.add_table(doc, headers, rows)
                    i = end_idx
                    continue
            
            # Detectar headings markdown
            if stripped.startswith('### '):
                if in_list and list_buffer:
                    self._flush_list(doc, list_buffer)
                    list_buffer = []
                    in_list = False
                self._add_formatted_paragraph(doc, stripped[4:], 'Heading 3')
                doc.add_paragraph()  # Espacio extra despues de heading
            elif stripped.startswith('## '):
                if in_list and list_buffer:
                    self._flush_list(doc, list_buffer)
                    list_buffer = []
                    in_list = False
                self._add_formatted_paragraph(doc, stripped[3:], 'Heading 2')
                doc.add_paragraph()  # Espacio extra despues de heading
            elif stripped.startswith('# '):
                if in_list and list_buffer:
                    self._flush_list(doc, list_buffer)
                    list_buffer = []
                    in_list = False
                self._add_formatted_paragraph(doc, stripped[2:], 'Heading 1')
                doc.add_paragraph()  # Espacio extra despues de heading
            # Detectar listas
            elif stripped.startswith('- ') or stripped.startswith('* '):
                in_list = True
                list_buffer.append(stripped[2:])
            # Línea vacía: flush list
            elif not stripped and in_list:
                self._flush_list(doc, list_buffer)
                list_buffer = []
                in_list = False
            # Texto normal
            elif stripped:
                if in_list and list_buffer:
                    self._flush_list(doc, list_buffer)
                    list_buffer = []
                    in_list = False
                self._add_formatted_paragraph(doc, stripped)
            
            i += 1
        
        # Flush remaining list
        if in_list and list_buffer:
            self._flush_list(doc, list_buffer)
        # Flush remaining code
        if in_code_block and code_buffer:
            self.add_code_block(doc, '\n'.join(code_buffer), language)

    def generate(self, metadata: Dict[str, Any], sections: List[Dict[str, Any]], output_path: str):
        """Genera el documento DOCX completo."""
        doc = Document()
        self.document = doc

        # Configurar estilos
        self.setup_styles(doc)
        
        # Configurar header/footer corporativos
        self.add_header_footer(doc, metadata)
        
        # Agregar borde de página
        self.add_page_border(doc)

        # Sección: Portada (Logo + Título + Metadatos)
        self.add_logo(doc)
        self.add_title(doc, metadata)
        self.add_metadata_section(doc, metadata)

        # Agregar cada sección
        for section in sections:
            sec_id = section.get('id', '')
            sec_title = section.get('title', '')
            sec_text = section.get('text', '')
            sec_table = section.get('table')
            
            # Verificar si la tabla debe incluirse
            include_table = section.get('include_table', True)

            if sec_id == 'portada':
                continue  # Ya procesada arriba
            if sec_id == 'indice':
                # Generar índice nativo de Word (TOC)
                self.add_toc(doc, sections)
                continue

            self.add_content_section(
                doc,
                sec_title,
                sec_text,
                sec_table if include_table else None
            )

        # Guardar
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path
